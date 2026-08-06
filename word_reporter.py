# -*- coding: utf-8 -*-
"""
word_reporter.py
v1.0 -- Local folder archiving + Word (.docx) auto-generation module.

Purpose:
    This module exposes a single core entry point, save_feedback_to_word(data: dict),
    which api_server.py's POST /api/save_word_report endpoint calls directly to do
    two things:

    1) Build a strict local directory tree:
           AI-Football-Feedback/student feedback report/
               |-- <realtime feedback folder> or <delayed feedback folder>  (level 1: test mode)
               |     |-- <school>-<class/group>/                            (level 2: school + class/group)
               |           |-- <student number>/                           (level 3: student number)

       Every level is created recursively with os.makedirs(..., exist_ok=True) if missing.

    2) Render a research-style Word (.docx) clinical report:
       Title ([student] - biomechanics realtime diagnosis) + metadata table
       (timestamp / score / T_impact) + optional key-frame image + four AI modules
       (overview / biomechanical_analysis / deductions table / magic_metaphor+action_plan),
       saved as "YYYY-MM-DD_HH-mm_<student>_<report>.docx". Absolute path is returned
       so api_server.py can hand it straight back to the frontend.

Robustness notes:
    - Base64 image decoding is defensive: whether the frontend sends a raw Base64 string,
      a standard "data:image/jpeg;base64,xxxx" data URI, or something missing/corrupted,
      decoding failures are swallowed silently -- the picture is simply skipped, and the
      rest of the report is still generated and saved normally.
    - Windows-illegal filename/foldername characters (backslash, slash, colon, asterisk,
      question mark, double quote, angle brackets, pipe) coming from free-text fields
      (school name, class/group name, student number) are all sanitized through
      sanitize_path_component() before touching the filesystem, so os.makedirs() /
      Document.save() never raise OSError and abort an otherwise successful save.
    - Module import time forces sys.stdout/sys.stderr to UTF-8 (errors='replace') so
      Windows's legacy GBK console code page never raises UnicodeEncodeError on CJK text
      or emoji in log lines (first line of defense). Terminal printing is additionally
      wrapped in _safe_print() (second line of defense): any UnicodeEncodeError that still
      slips through is caught and the message is safely re-encoded, so a logging hiccup
      never aborts an otherwise successful report generation.
"""

from __future__ import annotations

import base64
import io
import os
import re
import sys
import time
from pathlib import Path
from typing import Optional

# --------------------------------------------------------------------------
# Windows console-encoding compatibility fix (first line of defense).
#
# On Windows, the default console code page is GBK (cp936). When this module
# is imported by api_server.py and runs inside a background thread, any
# print()/log line containing a character outside the GBK charset (e.g. the
# report status is embedded together with debug text elsewhere in the
# pipeline) would raise UnicodeEncodeError and could abort the background
# archiving thread. Force stdout/stderr to UTF-8 with errors='replace' here
# so this never happens, regardless of what the OS console code page is.
# --------------------------------------------------------------------------
if sys.stdout.encoding is None or sys.stdout.encoding.lower() != "utf-8":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    except (AttributeError, ValueError):
        pass
if sys.stderr.encoding is None or sys.stderr.encoding.lower() != "utf-8":
    try:
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)
    except (AttributeError, ValueError):
        pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------
# All user-facing Chinese text is defined here via \uXXXX escapes instead of
# literal CJK characters. This keeps the source file 100% ASCII on disk,
# sidestepping any encoding round-trip issues in the local toolchain while
# still producing perfectly correct Chinese strings at runtime (Python
# decodes \uXXXX escapes the same way regardless of the file's own encoding).
# --------------------------------------------------------------------------

_MODE_REALTIME_LABEL = "\u5b9e\u65f6\u53cd\u9988"  # ????
_MODE_DELAYED_LABEL = "\u5ef6\u65f6\u53cd\u9988"  # ????

_FONT_HEADING_EASTASIA = "\u9ed1\u4f53"  # ??
_FONT_BODY_EASTASIA = "\u5fae\u8f6f\u96c5\u9ed1"  # ????

_FALLBACK_UNNAMED = "\u672a\u547d\u540d"  # ???
_FALLBACK_CLASS_FOLDER = "\u672a\u5206\u7c7b\u73ed\u7ea7"  # ?????
_FALLBACK_STUDENT_FOLDER = "\u672a\u586b\u5199\u5b66\u53f7"  # ?????
_FALLBACK_SCHOOL_TEXT = "\u672a\u8bbe\u7f6e\u5b66\u6821"  # ?????
_FALLBACK_CLASSGROUP_TEXT = "\u672a\u8bbe\u7f6e\u73ed\u7ea7"  # ?????
_FALLBACK_STUDENT_NUM_TEXT = "\u672a\u586b\u5199\u7f16\u53f7"  # ?????

_NO_SCORE_TEXT = "\u6682\u65e0\u8bc4\u5206"  # ????
_NO_DATA_TEXT = "\u6682\u65e0\u6570\u636e"  # ????

_LABEL_TIMESTAMP = "\u6d4b\u8bd5\u65e5\u671f\u4e0e\u65f6\u95f4\u6233"  # ????????
_LABEL_SCHOOL_CLASS = "\u5b66\u6821\u73ed\u7ea7"  # ????
_LABEL_STUDENT_NUM = "\u5b66\u751f\u7f16\u53f7"  # ????
_LABEL_SCORE = "\u53d1\u529b\u7efc\u5408\u8bc4\u5206"  # ??????
_LABEL_SAMPLE_COUNT = "\u6709\u6548\u91c7\u6837\u6b21\u6570"  # ??????

_UNIT_SCORE_SUFFIX = " \u5206"  # " ?"
_UNIT_TIMES_SUFFIX = " \u6b21"  # " ?"

_TITLE_SUFFIX = (
    " - \u8db3\u7403\u751f\u7269\u529b\u5b66\u5b9e\u65f6\u8bca\u65ad\u62a5\u544a"
)  # " - 足球生物力学实时诊断报告"
_SUBTITLE_SUFFIX = " - \u7cfb\u7edf\u81ea\u52a8\u5f52\u6863\u751f\u6210"  # " - 系统自动归档生成"

_IMAGE_CAPTION = (
    "\u4e0a\u56fe\uff1a\u51fb\u7403\u77ac\u95f4\u751f\u7269\u529b\u5b66\u5173\u952e\u5e27\u6807\u6ce8\u56fe"
    "\uff08\u9acb-\u819d-\u8e1d\u52a8\u529b\u94fe\u77e2\u91cf\uff09"
)

# 四维诊断模块标题
_HEADING_OVERVIEW = "\U0001f4c8 \u7efc\u5408\u4f53\u6001\u8bc4\u4ef7"  # 📈 综合体态评价
_HEADING_BIOMECH = "\U0001fa7a \u52a8\u529b\u94fe\u75c5\u7406\u5206\u6790"  # 🩺 动力链病理分析
_HEADING_DEDUCTIONS = "\U0001f4ca \u91cf\u5316\u6263\u5206\u660e\u7ec6"  # 📊 量化扣分明细
_HEADING_PRESCRIPTION = (
    "\U0001fa84 \u6559\u7ec3\u5904\u65b9\u4e0e\u8bad\u7ec3\u8ba1\u5212"
)  # 🪄 教练处方与训练计划

_NO_TEXT_FALLBACK = "\u6682\u65e0\u6570\u636e"  # 暂无数据

_FILENAME_SUFFIX = "\u8bca\u65ad\u5904\u65b9"  # 诊断处方

_LABEL_BRACKET_LEFT = "\u3010"  # 【
_LABEL_BRACKET_RIGHT = "\u3011"  # 】

_LABEL_T_IMPACT = "\u89e6\u7403\u5e27 T_impact"  # 触球帧 T_impact
_LABEL_TOTAL_SCORE = "\u7efc\u5408\u603b\u5206"  # 综合总分

_DEDUCTION_COL_NAME = "\u9519\u8bef\u540d\u79f0"  # 错误名称
_DEDUCTION_COL_VALUE = "\u5b9e\u6d4b\u6570\u503c"  # 实测数值
_DEDUCTION_COL_PENALTY = "\u6263\u5206"  # 扣分
_DEDUCTION_COL_REASON = "\u539f\u56e0"  # 原因
_DEDUCTION_COL_PROVENANCE = "\u6570\u636e\u8840\u7edf"  # 数据血统
_DEDUCTION_EMPTY = (
    "\u672c\u6b21\u65e0\u663e\u8457\u6263\u5206\u9879\uff0c"
    "\u52a8\u4f5c\u6574\u4f53\u5904\u4e8e\u53ef\u63a5\u53d7\u533a\u95f4\u3002"
)

_HEADING_PROVENANCE = (
    "\u6570\u636e\u8840\u7edf\uff08PROVENANCE\uff09\u900f\u660e\u5316"
)  # 数据血统（PROVENANCE）透明化

_PROVENANCE_TIER_LABEL_ZH = {
    "MEASURED": "MEASURED \u5b9e\u6d4b",  # 实测
    "CALIBRATED": "CALIBRATED \u6807\u5b9a",  # 标定
    "ESTIMATED": "ESTIMATED \u4f30\u7b97",  # 估算
}

_PROVENANCE_TIER_NOTE_ZH = {
    "MEASURED": "\u57fa\u4e8e\u7269\u7406\u50cf\u7d20\u5b9e\u6d4b\uff0c\u7f6e\u4fe1\u5ea6\u6781\u9ad8",
    "CALIBRATED": (
        "\u57fa\u4e8e\u89e3\u5256\u5b66/\u73af\u5883\u6807\u5b9a\u63a8\u7b97\uff0c"
        "\u7f6e\u4fe1\u5ea6\u9ad8"
    ),
    "ESTIMATED": (
        "\u57fa\u4e8e\u6df1\u5ea6\u5b66\u4e60 3D \u4f30\u7b97\uff0c"
        "\u4ec5\u4f9b\u53c2\u8003"
    ),
}

_METRIC_LABEL_ZH = {
    "ankle_rigidity": "\u811a\u8e1d\u9501\u7d27",
    "distance_cm": "\u652f\u6491\u811a\u7ad9\u4f4d",
    "toe_angle": "\u652f\u6491\u811a\u5c16\u65b9\u5411",
    "max_folding_angle": "\u6446\u52a8\u817f\u6298\u53e0",
    "impact_knee_angle": "\u89e6\u7403\u819d\u76d6",
    "support_knee_angle": "\u652f\u6491\u819d\u76d6",
    "hip_torsion_angle": "\u8f6c\u9acb",
    "whipping_velocity": "\u6446\u817f\u901f\u5ea6",
    "trunk_lean_angle": "\u8eaf\u5e72\u503e\u89d2",
}

_PRESCRIPTION_MAGIC_PREFIX = "\u3010\u5177\u8eab\u9690\u55bb\u5904\u65b9\u3011"
_PRESCRIPTION_ACTION_PREFIX = "\u3010\u4e0b\u4e00\u6b65\u8bad\u7ec3\u6307\u4ee4\u3011"


def _safe_print(message: str) -> None:
    """Print a log line (second line of defense): even though the module-level
    stdout/stderr UTF-8 reconfiguration above should already make this a
    non-issue, this still catches any stray UnicodeEncodeError and degrades
    the message to a safely re-encodable form instead of letting a logging
    hiccup abort an otherwise successful report generation.
    """
    try:
        print(message)
    except UnicodeEncodeError:
        encoding = getattr(sys.stdout, "encoding", None) or "utf-8"
        try:
            print(message.encode(encoding, errors="replace").decode(encoding, errors="replace"))
        except Exception:
            pass


# --------------------------------------------------------------------------
# Step 0: base paths & constants
# --------------------------------------------------------------------------

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Root archive folder is fixed to "<project root>/student feedback report/"
REPORT_ROOT_DIR = os.path.join(SCRIPT_DIR, "student feedback report")

# Level-1 subfolder: test mode -> Chinese folder name
MODE_FOLDER_NAME = {
    "realtime": _MODE_REALTIME_LABEL,
    "delayed": _MODE_DELAYED_LABEL,
}

# Target width (inches) when inserting the picture into the Word page: 5.5in
# renders crisp and centered on both A4 and Letter without overflowing margins.
IMAGE_WIDTH_INCHES = 5.5

# Windows-illegal filename/foldername characters: \ / : * ? " < > |
# plus all ASCII control characters (0x00-0x1F, e.g. stray \n / \t / \r that
# might leak in from free-text frontend fields), which Windows also rejects.
_ILLEGAL_CHARS_PATTERN = re.compile(r'[\\/:*?"<>|\x00-\x1f]')


def _set_run_east_asian_font(run, font_name: str) -> None:
    """Explicitly set the East-Asian font of a run so CJK glyphs actually
    render with the requested typeface (Word stores "Western" and "East
    Asian" fonts as two separate properties -- setting run.font.name alone
    only affects the Western half and leaves CJK text on the system default).
    """
    run.font.name = font_name
    run_properties = run._element.get_or_add_rPr()
    font_element = run_properties.find(qn("w:rFonts"))
    if font_element is None:
        font_element = run_properties.makeelement(qn("w:rFonts"), {})
        run_properties.append(font_element)
    font_element.set(qn("w:eastAsia"), font_name)


# --------------------------------------------------------------------------
# Step 1: path sanitizing & directory tree construction
# --------------------------------------------------------------------------


def sanitize_path_component(raw: Optional[str], fallback: str = _FALLBACK_UNNAMED) -> str:
    """Turn an arbitrary free-text string into a safe Windows folder/file
    name segment.

    - Blank/None falls back to `fallback`.
    - Windows-illegal characters are replaced with a hyphen "-".
    - Collapse consecutive hyphens produced by the substitution above, and
      strip leading/trailing dots or spaces (Windows disallows folder names
      ending in a dot or a space).
    """
    text = (raw or "").strip()
    if not text:
        text = fallback
    text = _ILLEGAL_CHARS_PATTERN.sub("-", text)
    text = re.sub(r"-{2,}", "-", text).strip(" .")
    return text or fallback


def build_target_directory(mode: str, school: str, class_group: str, student_number: str) -> Path:
    """Build (and recursively create) the full 3-level directory tree:
    level 1 test mode -> level 2 "<school>-<class/group>" -> level 3 student number.
    Returns the level-3 directory path.
    """
    mode_folder = MODE_FOLDER_NAME.get(mode, _MODE_REALTIME_LABEL)

    school_clean = (school or "").strip()
    class_group_clean = (class_group or "").strip()
    if school_clean and class_group_clean:
        school_class_raw = f"{school_clean}-{class_group_clean}"
    else:
        school_class_raw = school_clean or class_group_clean
    school_class_folder = sanitize_path_component(school_class_raw, _FALLBACK_CLASS_FOLDER)

    student_folder = sanitize_path_component(student_number, _FALLBACK_STUDENT_FOLDER)

    target_dir = Path(REPORT_ROOT_DIR) / mode_folder / school_class_folder / student_folder
    os.makedirs(target_dir, exist_ok=True)
    return target_dir


# --------------------------------------------------------------------------
# Step 2: defensive Base64 image decoding
# --------------------------------------------------------------------------


def _decode_base64_image_to_stream(image_base64: Optional[str]) -> Optional[io.BytesIO]:
    """Safely decode a Base64 image string (either a raw Base64 payload, or a
    standard "data:image/...;base64,xxxx" data URI) into an in-memory
    io.BytesIO stream.

    Any failure (missing/empty/malformed/corrupted input) is caught and
    returns None -- callers treat that as "skip the picture", never letting a
    bad image field abort the entire text report from being generated.
    """
    if not image_base64 or not isinstance(image_base64, str):
        return None
    try:
        payload = image_base64.strip()
        if payload.lower().startswith("data:") and "," in payload:
            payload = payload.split(",", 1)[1]
        raw_bytes = base64.b64decode(payload, validate=False)
        if not raw_bytes:
            return None
        return io.BytesIO(raw_bytes)
    except Exception as exc:  # noqa: BLE001 - image decoding must never abort the whole report
        _safe_print(f"[word_reporter] warning: base64 image decode failed, skipping image. reason: {exc}")
        return None


# --------------------------------------------------------------------------
# Step 3: Word (.docx) layout rendering
# --------------------------------------------------------------------------


def _safe_display(value, fallback: str = _NO_DATA_TEXT) -> str:
    """Null-safe display helper for Word cells: None / blank / NaN → fallback."""
    if value is None:
        return fallback
    if isinstance(value, float):
        try:
            import math

            if math.isnan(value) or math.isinf(value):
                return fallback
        except Exception:  # noqa: BLE001
            return fallback
    text = str(value).strip()
    if not text or text.lower() in ("none", "null", "undefined", "nan"):
        return fallback
    return text


def _safe_score_display(score) -> str:
    """Score cell: empty →「暂无评分」; otherwise 「{n} 分」."""
    if score is None or score == "":
        return _NO_SCORE_TEXT
    try:
        numeric = float(score)
        if numeric != numeric:  # NaN
            return _NO_SCORE_TEXT
        # Prefer integer display when whole number
        if numeric == int(numeric):
            return f"{int(numeric)}{_UNIT_SCORE_SUFFIX}"
        return f"{numeric:.1f}{_UNIT_SCORE_SUFFIX}"
    except (TypeError, ValueError):
        return _NO_SCORE_TEXT


def _safe_count_display(count) -> str:
    """Sample-count cell: empty →「暂无数据」."""
    if count is None or count == "":
        return _NO_DATA_TEXT
    try:
        numeric = float(count)
        if numeric != numeric:
            return _NO_DATA_TEXT
        return f"{int(round(numeric))}{_UNIT_TIMES_SUFFIX}"
    except (TypeError, ValueError):
        return _NO_DATA_TEXT


def _safe_error_codes_display(data: dict) -> str:
    """Biomechanical / error-code list → readable Chinese, or「暂无数据」."""
    if not isinstance(data, dict):
        return _NO_DATA_TEXT
    raw = (
        data.get("biomechanicalErrors")
        or data.get("biomechanical_errors")
        or data.get("errorCodes")
        or data.get("primaryErrorCode")
        or data.get("primary_error_code")
    )
    if raw is None or raw == "":
        return _NO_DATA_TEXT
    if isinstance(raw, (list, tuple, set)):
        parts = [_safe_display(item, "") for item in raw]
        parts = [p for p in parts if p]
        return "、".join(parts) if parts else _NO_DATA_TEXT
    return _safe_display(raw, _NO_DATA_TEXT)


def _add_metadata_table(document: Document, rows: list[tuple[str, str]]) -> None:
    """Insert a two-column metadata table below the title: bold label on the
    left, corresponding value on the right.
    """
    safe_rows = [
        (_safe_display(label, ""), _safe_display(value, _NO_DATA_TEXT))
        for label, value in (rows or [])
    ]
    if not safe_rows:
        safe_rows = [(_NO_DATA_TEXT, _NO_DATA_TEXT)]

    table = document.add_table(rows=len(safe_rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for row_index, (label, value) in enumerate(safe_rows):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)

        label_cell.text = ""
        label_paragraph = label_cell.paragraphs[0]
        label_run = label_paragraph.add_run(
            f"{_LABEL_BRACKET_LEFT}{label}{_LABEL_BRACKET_RIGHT}"
        )
        label_run.bold = True
        label_run.font.size = Pt(11)
        _set_run_east_asian_font(label_run, _FONT_BODY_EASTASIA)

        value_cell.text = ""
        value_paragraph = value_cell.paragraphs[0]
        value_run = value_paragraph.add_run(_safe_display(value, _NO_DATA_TEXT))
        value_run.font.size = Pt(11)
        _set_run_east_asian_font(value_run, _FONT_BODY_EASTASIA)


def _add_section_heading(document: Document, text: str, level: int = 2) -> None:
    """Insert a Heading-style section title (科研体检报告模块标题)."""
    try:
        heading = document.add_heading(_safe_display(text, _NO_DATA_TEXT), level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            _set_run_east_asian_font(run, _FONT_HEADING_EASTASIA)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
    except Exception:
        # Heading 样式异常时退化为加粗段落，绝不中断报告
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(_safe_display(text, _NO_DATA_TEXT))
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_run_east_asian_font(run, _FONT_HEADING_EASTASIA)


def _add_body_paragraph(document: Document, text: str) -> None:
    """Insert a body paragraph with comfortable spacing, ready to print."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(_safe_display(text, _NO_TEXT_FALLBACK))
    run.font.size = Pt(12)
    _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)


def _add_emphasis_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = True,
    color: Optional[RGBColor] = None,
) -> None:
    """Insert an emphasized body paragraph (教练处方高亮)."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(_safe_display(text, _NO_TEXT_FALLBACK))
    run.bold = bool(bold)
    run.font.size = Pt(12)
    run.font.color.rgb = color or RGBColor(0xC0, 0x56, 0x21)
    _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)


_LABEL_ERROR_CODES = "\u751f\u7269\u529b\u5b66\u9519\u8bef\u7801"  # 生物力学错误码


def _extract_score_detail(data: dict) -> dict:
    """Safely pull score_detail / scoreDetail from the archive payload."""
    if not isinstance(data, dict):
        return {}
    detail = data.get("scoreDetail") or data.get("score_detail") or {}
    return detail if isinstance(detail, dict) else {}


def _extract_ai_dimensions(data: dict) -> dict:
    """Safe-extract four AI report fields with legacy fallbacks.

    Never raises; missing keys become「暂无数据」.
    """
    data = data if isinstance(data, dict) else {}

    def _pick(*keys: str) -> str:
        for key in keys:
            raw = data.get(key)
            if raw is None:
                continue
            text = str(raw).strip()
            if text and text.lower() not in ("none", "null", "undefined", "nan"):
                return text
        return _NO_TEXT_FALLBACK

    return {
        "overview": _pick("overview", "clinical_echo", "clinicalEcho"),
        "biomechanical_analysis": _pick(
            "biomechanical_analysis", "painPoint", "comment"
        ),
        "magic_metaphor": _pick(
            "magic_metaphor", "correction_metaphor", "painPoint"
        ),
        "action_plan": _pick(
            "action_plan", "praise_encouragement", "prescription"
        ),
    }


def _resolve_provenance_tier_label(entry_or_row: dict) -> str:
    """Normalize provenance_tier for Word display."""
    if not isinstance(entry_or_row, dict):
        return _NO_DATA_TEXT
    raw = (
        entry_or_row.get("provenance_tier")
        or entry_or_row.get("provenanceTier")
        or ""
    )
    tier = str(raw).strip().upper()
    if tier in _PROVENANCE_TIER_LABEL_ZH:
        return _PROVENANCE_TIER_LABEL_ZH[tier]
    # 小写 provenance 回退
    prov = str(entry_or_row.get("provenance") or "").strip().lower()
    if prov == "measured":
        return _PROVENANCE_TIER_LABEL_ZH["MEASURED"]
    if prov == "calibrated":
        return _PROVENANCE_TIER_LABEL_ZH["CALIBRATED"]
    if prov in ("estimated", "default", "missing"):
        return _PROVENANCE_TIER_LABEL_ZH["ESTIMATED"]
    return _NO_DATA_TEXT


def _format_measured_value(row: dict) -> str:
    """Format deduction measured_value + unit for the table cell."""
    if not isinstance(row, dict):
        return _NO_DATA_TEXT
    value = row.get("measured_value")
    if value is None:
        value = row.get("value")
    if value is None and row.get("metric_key") == "ankle_rigidity":
        value = row.get("variance")
    if value is None:
        return _NO_DATA_TEXT
    unit = str(row.get("unit") or "").strip()
    metric_key = str(row.get("metric_key") or row.get("metric") or "")
    if not unit and metric_key == "ankle_rigidity":
        unit = "variance"
    elif not unit and metric_key in {
        "impact_knee_angle",
        "support_knee_angle",
        "max_folding_angle",
        "toe_angle",
        "hip_torsion_angle",
        "trunk_lean_angle",
    }:
        unit = "deg"
    elif not unit and metric_key == "distance_cm":
        unit = "cm"
    elif not unit and metric_key == "whipping_velocity":
        unit = "deg/s"
    try:
        num = float(value)
        if unit in {"deg", "°", "degree"}:
            return f"{num:.1f}°"
        if unit in {"cm"}:
            return f"{num:.1f}cm"
        if unit in {"ratio", "×肩宽", "x肩宽"}:
            return f"{num:.2f}×肩宽"
        if unit in {"variance", "σ²"}:
            return f"σ² {num:.2f}"
        if unit in {"deg/s", "°/s"}:
            return f"{num:.1f}°/s"
        if unit:
            return f"{num:.2f}{unit}"
        return f"{num:.2f}"
    except (TypeError, ValueError):
        return _safe_display(value, _NO_DATA_TEXT)


def _deduction_metric_label(row: dict) -> str:
    """Human-readable error name for a deduction row."""
    if not isinstance(row, dict):
        return _NO_DATA_TEXT
    key = row.get("metric_key") or row.get("metric") or row.get("key")
    if isinstance(key, str) and key in _METRIC_LABEL_ZH:
        return _METRIC_LABEL_ZH[key]
    if isinstance(key, str) and key.strip():
        return key.strip()
    reason = str(row.get("reason") or "").strip()
    return reason[:12] if reason else _NO_DATA_TEXT


def _collect_deduction_rows(data: dict) -> list[dict]:
    """Collect deduction rows from score_detail; tolerate missing/malformed data.

    若归档快照丢失 deductions，则从 indicators.penalty>0 重建表格行。
    """
    detail = _extract_score_detail(data)
    indicators = (
        detail.get("indicators") if isinstance(detail.get("indicators"), dict) else {}
    )
    raw = detail.get("deductions") or data.get("deductions") or []
    rows: list[dict] = []
    if isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict):
                rows.append(item)
            elif isinstance(item, str) and item.strip():
                rows.append(
                    {
                        "metric_key": "",
                        "measured_value": None,
                        "unit": "",
                        "penalty": "",
                        "reason": item.strip(),
                    }
                )

    if not rows:
        for key, item in indicators.items():
            if not isinstance(item, dict):
                continue
            try:
                penalty = float(item.get("penalty") or 0.0)
            except (TypeError, ValueError):
                penalty = 0.0
            if penalty <= 0:
                continue
            value = item.get("value")
            if value is None:
                value = item.get("variance")
            rows.append(
                {
                    "metric_key": key,
                    "measured_value": value,
                    "unit": item.get("unit") or "",
                    "penalty": penalty,
                    "reason": f"{_METRIC_LABEL_ZH.get(key, key)}偏离理想（{item.get('status') or '偏离'}）",
                    "status": item.get("status"),
                    "provenance_tier": item.get("provenance_tier")
                    or item.get("provenanceTier"),
                    "provenance": item.get("provenance"),
                }
            )

    # 补全扣分行血统：优先行内字段，否则从 indicators 回填
    for row in rows:
        if not isinstance(row, dict):
            continue
        if row.get("provenance_tier") or row.get("provenanceTier"):
            continue
        key = str(row.get("metric_key") or "")
        src = indicators.get(key) if key else None
        if isinstance(src, dict):
            row["provenance_tier"] = src.get("provenance_tier") or src.get(
                "provenanceTier"
            )
            row.setdefault("provenance", src.get("provenance"))
    rows.sort(key=lambda r: -float(r.get("penalty") or 0.0))
    return rows


def _text_has_measurement(text: str) -> bool:
    """Heuristic: data-driven copy should contain at least one digit."""
    return bool(re.search(r"\d", text or ""))


def _synthesize_ai_from_score_detail(data: dict) -> dict:
    """When AI four fields lack numbers, rebuild from scoreDetail via llm_agent."""
    detail = _extract_score_detail(data)
    if not isinstance(detail, dict) or not detail:
        return {}
    try:
        import llm_agent as la

        detail_for_llm = dict(detail)
        if detail_for_llm.get("TotalScore") is None and data.get("score") is not None:
            detail_for_llm["TotalScore"] = data.get("score")
        synth = la._depth_fallback_report({"score_detail": detail_for_llm})
        return {
            "overview": synth.get("overview") or _NO_TEXT_FALLBACK,
            "biomechanical_analysis": synth.get("biomechanical_analysis") or _NO_TEXT_FALLBACK,
            "magic_metaphor": synth.get("magic_metaphor") or _NO_TEXT_FALLBACK,
            "action_plan": synth.get("action_plan") or _NO_TEXT_FALLBACK,
        }
    except Exception as exc:  # noqa: BLE001
        _safe_print(f"[word_reporter] warning: synthesize AI from scoreDetail failed: {exc}")
        return {}


def _resolve_ai_dimensions(data: dict) -> dict:
    """Prefer payload four fields; if they lack measured numbers, synthesize."""
    ai = _extract_ai_dimensions(data)
    detail = _extract_score_detail(data)
    has_metrics = bool(
        (isinstance(detail.get("indicators"), dict) and detail.get("indicators"))
        or (isinstance(detail.get("deductions"), list) and detail.get("deductions"))
    )
    if not has_metrics:
        return ai

    needs_synth = (
        not _text_has_measurement(ai.get("overview", ""))
        or not _text_has_measurement(ai.get("biomechanical_analysis", ""))
        or ai.get("overview") == _NO_TEXT_FALLBACK
        or ai.get("biomechanical_analysis") == _NO_TEXT_FALLBACK
    )
    if not needs_synth:
        return ai

    synth = _synthesize_ai_from_score_detail(data)
    if not synth:
        return ai
    for key in ("overview", "biomechanical_analysis", "magic_metaphor", "action_plan"):
        if (
            not _text_has_measurement(ai.get(key, ""))
            or ai.get(key) in ("", _NO_TEXT_FALLBACK)
        ):
            ai[key] = synth.get(key) or ai.get(key) or _NO_TEXT_FALLBACK
    return ai


def _add_deductions_table(document: Document, data: dict) -> None:
    """Insert a research-style deductions table (错误名称 / 实测 / 血统 / 扣分 / 原因)."""
    rows = _collect_deduction_rows(data)
    if not rows:
        _add_body_paragraph(document, _DEDUCTION_EMPTY)
        return

    table = document.add_table(rows=1 + len(rows), cols=5)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    table.autofit = True

    headers = (
        _DEDUCTION_COL_NAME,
        _DEDUCTION_COL_VALUE,
        _DEDUCTION_COL_PROVENANCE,
        _DEDUCTION_COL_PENALTY,
        _DEDUCTION_COL_REASON,
    )
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)

    for row_idx, row in enumerate(rows, start=1):
        penalty = row.get("penalty")
        try:
            penalty_text = f"{float(penalty):.1f}" if penalty is not None and penalty != "" else _NO_DATA_TEXT
        except (TypeError, ValueError):
            penalty_text = _safe_display(penalty, _NO_DATA_TEXT)

        values = (
            _deduction_metric_label(row),
            _format_measured_value(row),
            _resolve_provenance_tier_label(row),
            penalty_text,
            _safe_display(row.get("reason"), _NO_DATA_TEXT),
        )
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(_safe_display(value, _NO_DATA_TEXT))
            run.font.size = Pt(10)
            _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)

    document.add_paragraph()


def _add_provenance_transparency_table(document: Document, data: dict) -> None:
    """List each indicator's provenance_tier for full-chain transparency."""
    detail = _extract_score_detail(data)
    indicators = (
        detail.get("indicators") if isinstance(detail.get("indicators"), dict) else {}
    )
    if not indicators:
        return

    # 图例说明
    legend_bits = [
        f"{_PROVENANCE_TIER_LABEL_ZH[k]}：{_PROVENANCE_TIER_NOTE_ZH[k]}"
        for k in ("MEASURED", "CALIBRATED", "ESTIMATED")
    ]
    _add_body_paragraph(document, " · ".join(legend_bits))

    keys = [
        k
        for k in (
            "distance_cm",
            "toe_angle",
            "max_folding_angle",
            "whipping_velocity",
            "impact_knee_angle",
            "ankle_rigidity",
            "support_knee_angle",
            "hip_torsion_angle",
            "trunk_lean_angle",
        )
        if k in indicators
    ]
    # 追加未在标准顺序中的键
    for k in indicators:
        if k not in keys:
            keys.append(k)
    if not keys:
        return

    table = document.add_table(rows=1 + len(keys), cols=3)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    table.autofit = True

    headers = (
        "\u6307\u6807",  # 指标
        _DEDUCTION_COL_VALUE,
        _DEDUCTION_COL_PROVENANCE,
    )
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)

    for row_idx, key in enumerate(keys, start=1):
        item = indicators.get(key) if isinstance(indicators.get(key), dict) else {}
        row_proxy = {
            "metric_key": key,
            "measured_value": item.get("value"),
            "unit": item.get("unit") or "",
            "variance": item.get("variance"),
            "provenance_tier": item.get("provenance_tier") or item.get("provenanceTier"),
            "provenance": item.get("provenance"),
        }
        values = (
            _METRIC_LABEL_ZH.get(key, key),
            _format_measured_value(row_proxy),
            _resolve_provenance_tier_label(row_proxy),
        )
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(_safe_display(value, _NO_DATA_TEXT))
            run.font.size = Pt(10)
            _set_run_east_asian_font(run, _FONT_BODY_EASTASIA)

    document.add_paragraph()


def _resolve_t_impact(data: dict) -> str:
    """Resolve T_impact index from top-level or score_detail."""
    detail = _extract_score_detail(data)
    for src in (data, detail):
        if not isinstance(src, dict):
            continue
        for key in ("t_impact", "tImpact", "t0_index"):
            raw = src.get(key)
            if raw is None or raw == "":
                continue
            try:
                return str(int(float(raw)))
            except (TypeError, ValueError):
                text = str(raw).strip()
                if text:
                    return text
    return _NO_DATA_TEXT


def _resolve_total_score_display(data: dict) -> str:
    """Prefer deterministic TotalScore from score_detail, else payload score."""
    detail = _extract_score_detail(data)
    for key in ("TotalScore", "total_score", "totalScore"):
        if detail.get(key) is not None:
            return _safe_score_display(detail.get(key))
    return _safe_score_display(data.get("score"))


def _build_document(data: dict, mode: str) -> Document:
    """Build the full Word document (科研体检报告质感四维排版).

    Every field is read via .get() / null-safe helpers so missing AI fields /
    deductions / score never raise and never write the literal string "None".
    """
    data = data if isinstance(data, dict) else {}

    school = _safe_display(data.get("school"), _FALLBACK_SCHOOL_TEXT)
    class_group = _safe_display(data.get("classGroup"), _FALLBACK_CLASSGROUP_TEXT)
    student_number = _safe_display(
        data.get("studentNumber") or data.get("studentId") or data.get("name"),
        _FALLBACK_STUDENT_NUM_TEXT,
    )
    score_display = _resolve_total_score_display(data)
    total_attempts_display = _safe_count_display(data.get("totalAttempts"))
    generated_at = _safe_display(
        data.get("generatedAt"), time.strftime("%Y-%m-%d %H:%M:%S")
    )
    t_impact_display = _resolve_t_impact(data)
    error_codes_display = _safe_error_codes_display(data)
    # 若前端仍传旧双字段/无数字模板，这里按 scoreDetail 重写为数据驱动四维文案
    ai = _resolve_ai_dimensions(data)

    document = Document()

    # --- Title: [学号] - 足球生物力学实时诊断报告 (Heading 1, 居中) ---
    title_text = f"{student_number}{_TITLE_SUFFIX}"
    try:
        title = document.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            _set_run_east_asian_font(run, _FONT_HEADING_EASTASIA)
    except Exception:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_run_east_asian_font(title_run, _FONT_HEADING_EASTASIA)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(
        f"{MODE_FOLDER_NAME.get(mode, _MODE_REALTIME_LABEL)}{_SUBTITLE_SUFFIX}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    _set_run_east_asian_font(subtitle_run, _FONT_BODY_EASTASIA)

    document.add_paragraph()

    # --- 测试基本信息 ---
    _add_metadata_table(
        document,
        rows=[
            (_LABEL_TIMESTAMP, generated_at),
            (_LABEL_SCHOOL_CLASS, f"{school} - {class_group}"),
            (_LABEL_STUDENT_NUM, student_number),
            (_LABEL_TOTAL_SCORE, score_display),
            (_LABEL_SCORE, score_display),
            (_LABEL_T_IMPACT, t_impact_display),
            (_LABEL_SAMPLE_COUNT, total_attempts_display),
            (_LABEL_ERROR_CODES, error_codes_display),
        ],
    )

    document.add_paragraph()

    # --- Biomechanics annotated key frame ---
    image_stream = _decode_base64_image_to_stream(data.get("impactFrameImage"))
    if image_stream is not None:
        try:
            image_paragraph = document.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_run = image_paragraph.add_run()
            image_run.add_picture(image_stream, width=Inches(IMAGE_WIDTH_INCHES))

            caption_paragraph = document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(_IMAGE_CAPTION)
            caption_run.font.size = Pt(9)
            caption_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            _set_run_east_asian_font(caption_run, _FONT_BODY_EASTASIA)
        except Exception as exc:  # noqa: BLE001
            _safe_print(f"[word_reporter] warning: insert image failed, skipping image. reason: {exc}")

    document.add_paragraph()

    # --- 模块 1：综合体态评价 ---
    _add_section_heading(document, _HEADING_OVERVIEW, level=2)
    _add_body_paragraph(document, ai["overview"])

    # --- 模块 2：动力链病理分析 ---
    _add_section_heading(document, _HEADING_BIOMECH, level=2)
    _add_body_paragraph(document, ai["biomechanical_analysis"])

    # --- 模块 3：量化扣分明细（表格）---
    _add_section_heading(document, _HEADING_DEDUCTIONS, level=2)
    try:
        _add_deductions_table(document, data)
    except Exception as exc:  # noqa: BLE001
        _safe_print(f"[word_reporter] warning: deductions table failed: {exc}")
        _add_body_paragraph(document, _DEDUCTION_EMPTY)

    # --- 模块 3b：数据血统（PROVENANCE）透明化 ---
    _add_section_heading(document, _HEADING_PROVENANCE, level=2)
    try:
        _add_provenance_transparency_table(document, data)
    except Exception as exc:  # noqa: BLE001
        _safe_print(f"[word_reporter] warning: provenance table failed: {exc}")

    # --- 模块 4：教练处方与训练计划（加粗 / 强调色）---
    _add_section_heading(document, _HEADING_PRESCRIPTION, level=2)
    _add_emphasis_paragraph(
        document,
        f"{_PRESCRIPTION_MAGIC_PREFIX}{ai['magic_metaphor']}",
        bold=True,
        color=RGBColor(0xC0, 0x56, 0x21),
    )
    _add_emphasis_paragraph(
        document,
        f"{_PRESCRIPTION_ACTION_PREFIX}{ai['action_plan']}",
        bold=True,
        color=RGBColor(0x1F, 0x6F, 0x4A),
    )

    return document


# --------------------------------------------------------------------------
# Step 4: public entry point
# --------------------------------------------------------------------------


def save_feedback_to_word(data: dict) -> dict:
    """Core entry point: given a data dict assembled by the frontend/backend,
    build the local directory tree and save the Word (.docx) report, then
    return a structured result dict.

    Expected fields in `data` (all defensively handled, none are strictly
    required to be present):
        mode            : "realtime" | "delayed"
        school          : str  -- school/institution name
        classGroup      : str  -- class / experiment group name
        studentNumber   : str  -- student number / ID
        score           : int | None
        totalAttempts   : int | None
        overview                 : str  -- 综合评价
        biomechanical_analysis   : str  -- 动力链病理分析
        magic_metaphor           : str  -- 具身隐喻处方
        action_plan              : str  -- 下一步训练指令
        painPoint / prescription : str  -- legacy fallbacks
        scoreDetail / score_detail: dict -- deductions / TotalScore / t_impact
        generatedAt     : str | None
        impactFrameImage: str | None  -- Base64 / data URI key frame image

    Returns:
        {"success": True, "path": "...", "directory": "...", "filename": "..."}
        or {"success": False, "error": "..."} on failure.
    """
    try:
        if not isinstance(data, dict):
            data = {}

        mode_raw = data.get("mode")
        mode = mode_raw if mode_raw in MODE_FOLDER_NAME else "realtime"
        school = _safe_display(data.get("school"), "")
        # Empty school should stay empty for folder builder (it has its own fallback)
        if school == _NO_DATA_TEXT:
            school = ""
        class_group_raw = data.get("classGroup")
        class_group = "" if class_group_raw is None else str(class_group_raw).strip()
        student_number_raw = (
            data.get("studentNumber") or data.get("studentId") or data.get("name") or ""
        )
        if student_number_raw is None:
            student_number_raw = ""

        target_dir = build_target_directory(mode, school, class_group, str(student_number_raw))

        try:
            document = _build_document(data, mode)
        except (KeyError, TypeError, ValueError, AttributeError) as build_exc:
            _safe_print(
                f"[word_reporter] error: build document failed "
                f"({type(build_exc).__name__}): {build_exc}"
            )
            return {
                "success": False,
                "error": f"报告生成失败，部分数据缺失: {build_exc}",
            }

        # File naming convention: YYYY-MM-DD_HH-mm_<student number>_<report>.docx
        timestamp_label = time.strftime("%Y-%m-%d_%H-%M")
        student_number_clean = sanitize_path_component(
            str(student_number_raw) if student_number_raw is not None else "",
            _FALLBACK_STUDENT_FOLDER,
        )
        filename = f"{timestamp_label}_{student_number_clean}_{_FILENAME_SUFFIX}.docx"

        full_path = target_dir / filename
        document.save(str(full_path))

        _safe_print(f"[word_reporter] saved Word report to: {full_path}")

        return {
            "success": True,
            "path": str(full_path.resolve()),
            "directory": str(target_dir.resolve()),
            "filename": filename,
        }
    except (KeyError, TypeError, ValueError, AttributeError) as exc:
        _safe_print(
            f"[word_reporter] error: save word report data missing "
            f"({type(exc).__name__}): {exc}"
        )
        return {"success": False, "error": f"报告生成失败，部分数据缺失: {exc}"}
    except Exception as exc:  # noqa: BLE001 - any failure must be reported, never crash the caller
        _safe_print(f"[word_reporter] error: save word report failed: {exc}")
        return {"success": False, "error": str(exc)}
